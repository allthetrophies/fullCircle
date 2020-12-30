#Library imports
import requests
import time
import re
import os
import wikipedia
import subprocess
import _tkinter
from docx import Document
from docx.shared import Inches
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import tkinter.font as tkFont

listOfWatchedFilms = []    #Value will be used for storing films for the watched list
listCastList = []          #Value will be used for storing the cast list of a given film
listFullCircle = []        #Value will be used for storing the list used for Full Circle
wikiResults = []           #Value will be used for storing the results from the wiki search
APIKEY = "ef9bd486181321a8c5dbd8d87432ecaf"   #Provided API Key that allows me to call TMDb










def onTabClickFunc(self):
    readOnlyText.configure(state="normal")
    readOnlyText.delete('1.0', END)
    if firstFilmTitleEntry.get():
        readOnlyText.insert(END, "First Film: ", "bold")
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, firstFilmTitleEntry.get() + " " + firstFilmYearEntry.get())
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, '\n')
    if secondFilmTitleEntry.get():
        readOnlyText.insert(END, "Second Film: ", "bold")
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, secondFilmTitleEntry.get() + " " + secondFilmYearEntry.get())
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, '\n')
    if openLineRadio.get() == 0:
        readOnlyText.insert(END, "Opening Line: ", "bold")
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, openingLineEntry.get())
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, '\n')
    else:
        readOnlyText.insert(END, "Opening Line: ", "bold")
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, "DISABLED")
        readOnlyText.insert(END, '\n')
        readOnlyText.insert(END, '\n')
    if len(listOfWatchedFilms) != 0:
        readOnlyText.insert(END, "Films I've Watched: ", "bold")
        readOnlyText.insert(END, '\n')
        for y in listOfWatchedFilms:
            if y != "":
                readOnlyText.insert(END, y)
                readOnlyText.insert(END, '\n')
    readOnlyText.configure(state="disabled")
    readOnlyText.update()










#Function for enabling the Opening Line field
def enableOpeningLineEntryFunc():
    openingLineEntry.configure(state="normal")    #Sets the entry field to normal, or enabled
    openingLineEntry.update()                     #Update the status of the entry field










#Function for disabling the Opening Line field
def disableOpeningLineEntryFunc():
    openingLineEntry.configure(state="disabled")    #Sets the entry field to disabled
    openingLineEntry.update()                       #Update the status of the entry field










#Function for adding in films to the watch list
def addToWatchListFunc():
    if filmWatchedEntry.index("end") == 0:
        messagebox.showerror("ERROR!", "Please make sure you are actually adding something to your Watched List")
    else:
        filmToAdd = filmWatchedEntry.get()      #Grab the film placed into the filmWatchedEntry field
        listOfWatchedFilms.append(filmToAdd)    #Append this film to the list of watched films
        filmWatchedEntry.delete(0, 'end')       #Empty out the filmWatchedEntry field

        #Read Only initialization for the Watch List
        filmWatchText.configure(state="normal")
        filmWatchText.delete('1.0', END)
        for y in listOfWatchedFilms:
            if y != "":
                filmWatchText.insert(END, y)
                filmWatchText.insert(END, '\n')
        filmWatchText.configure(state="disabled")
        filmWatchText.update()










#Function for obtaining a films cast list and formatting it
def castListFunc(filmTitle, filmYear):
    filmCastList = []         #Dynamic list used to hold cast members of entered film
    finalFilmCastList = []    #Dynamic list used to hold cast members with added format for output
    filmCharacter = ""        #Variable for storing the character that a given actor played

    #The film title without excess space and non-alphanumeric characters
    reducedFilmTitle = re.sub('[^A-Za-z0-9]+', '', filmTitle).lower()

    #Generated URL that allows us to grab the id of the movie provided by user input
    filmID_URL = "https://api.themoviedb.org/3/search/movie?api_key=" + APIKEY + "&query=" + filmTitle + "&year=" + filmYear

    #Takes the provided JSON and stores it into a new variable
    userResponse = requests.get(filmID_URL)
    IDGrab = userResponse.json()

    #Grabs the film ID of the film requested by the user
    for x in IDGrab['results']:
        if str(x['release_date'][0:4]) == filmYear and re.sub('[^A-Za-z0-9]+', '', str(x['title'])).lower() == reducedFilmTitle:
            filmID = str(x['id'])

    #API call that grabs the cast of the movie provided by user input, the call is accomplished with the film ID
    filmCastURL = "https://api.themoviedb.org/3/movie/" + filmID + "/credits?api_key=" + APIKEY

    #Take the JSON-ed cast and convert it into something easily readable for Python code
    castResponse = requests.get(filmCastURL)
    filmCastGrab = castResponse.json()

    #Store each cast member of the user's film into the dynamic list
    for x in filmCastGrab['cast']:
        filmCastList.append(x['name'])

    #Goes through every actor/actress in the film and pulls the character they played
    for x in set(filmCastList):
        actorURL = "http://api.tmdb.org/3/search/person?api_key=" + APIKEY + "&query=" + x

        #Take the JSON-ed actors and convert it into something easily readable for Python code
        actorResponse = requests.get(actorURL)
        actorGrab = actorResponse.json()

        #Hokey, don't replicate. This pulls the first ID present and then exits the loop
        for y in actorGrab['results']:
            actorID = str(y['id'])
            break

        #Obtain the URL for the characters played by the given actor/actress
        characterURL = "https://api.themoviedb.org/3/person/" + actorID + "/movie_credits?api_key=" + APIKEY

        #Take the JSON-ed characters and convert it into something easily readable for Python code
        characterResponse = requests.get(characterURL)
        characterGrab = characterResponse.json()

        #For every cast tag in the character list, find the one that matches the film title/release date passed in
        for y in characterGrab['cast']:
            if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedFilmTitle and str(y['release_date'][0:4]) == filmYear:
                filmCharacter = y['character']

        #If empty string, replace with error message. To be properly handled/fixed in the future
        if filmCharacter == "":
            filmCharacter = "**UNABLE TO PULL CHARACTER DATA**"
        
        #Add a formatted cast message to the cast list that will be returned at the end of the function
        finalFilmCastList.append(x + " plays " + filmCharacter + " in " + filmTitle + " (" + filmYear + ")")

        filmCharacter = ""    #Set the film character string back to empty

    return(finalFilmCastList)    #Return the cast list of the film that was passed in










#Function for obtaining Full Circle instances between the passed in film and previously covered films on the show
def fullCircleFunc(firstFilmTitle, firstFilmYear):
    firstFilmCastList = []       #Dynamic list used to hold cast members of entered film
    comparedFilmCastList = []    #Dynamic list used to hold cast members of film compared with cast of firstFilmCastList
    finalActorList = []          #Final list to hold all actors/actresses that fit the requirements
    comparedFilmTitle = ""       #Placeholder variable for the read-in film's title
    comparedFilmYear = ""        #Placeholder variable for the read-in film's year
    iterator = 0                 #Incremental variable for while loop below
    firstCharacter = ""          #First stored character role from the list intersection
    secondCharacter = ""         #Second stored character role from the list intersection

    #The first film title without excess space and non-alphanumeric characters
    reducedFirstFilmTitle = re.sub('[^A-Za-z0-9]+', '', firstFilmTitle).lower()

    #First generated URL that allows us to grab the id of the movie provided by user input
    firstFilmID_URL = "https://api.themoviedb.org/3/search/movie?api_key=" + APIKEY + "&query=" + firstFilmTitle + "&year=" + firstFilmYear

    #Takes the provided JSON and stores it into a new variable
    userResponse = requests.get(firstFilmID_URL)
    IDGrab = userResponse.json()

    #Grabs the film ID of the film requested by the user
    for x in IDGrab['results']:
        if str(x['release_date'][0:4]) == firstFilmYear and re.sub('[^A-Za-z0-9]+', '', str(x['title'])).lower() == reducedFirstFilmTitle:
            firstFilmID = str(x['id'])

    #API call that grabs the cast of the movie provided by user input, the call is accomplished with the film ID
    firstFilmCastURL = "https://api.themoviedb.org/3/movie/" + firstFilmID + "/credits?api_key=" + APIKEY

    #Take the JSON-ed cast and convert it into something easily readable for Python code
    castResponse = requests.get(firstFilmCastURL)
    firstFilmCastGrab = castResponse.json()

    #Store each cast member of the user's film into the dynamic list
    for x in firstFilmCastGrab['cast']:
        firstFilmCastList.append(x['name'])

    #Open up and read in the text file with all of the films covered by the They Remade It podcast
    filmListIn = open("../../textFiles/theyRemadeItList.txt", "r")
    filmCompareList = filmListIn.read().splitlines()

    #So, while the incremented variable is less than the length of the text file, this while loop will run through and make an API call for each
    #film in the text file, grab it's ID, make a call with the ID to grab the cast list of the film, compare the cast with that of the movie provided
    #through user input, and store any actors/actresses that show up in both into a new dynamic list
    while iterator < len(filmCompareList):
        comparedFilmTitle = filmCompareList[iterator][0:len(filmCompareList[iterator])-4]
        comparedFilmYear = filmCompareList[iterator][len(filmCompareList[iterator])-4:len(filmCompareList[iterator])]

        reducedComparedFilmTitle = re.sub('[^A-Za-z0-9]+', '', comparedFilmTitle).lower()

        #The URL for the film being compared to the original entry
        filmCompareURL = "https://api.themoviedb.org/3/search/movie?api_key=" + APIKEY + "&query=" + comparedFilmTitle + "&year=" + comparedFilmYear

        #Convert the JSON from the call into a usable film ID
        compareResponse = requests.get(filmCompareURL)
        compareGrab = compareResponse.json()

        #Used to grab the film ID listed in the pull request
        for x in compareGrab['results']:
            if str(x['release_date'][0:4]) == comparedFilmYear and re.sub('[^A-Za-z0-9]+', '', str(x['title'])).lower() == reducedComparedFilmTitle:
                filmCompareID = str(x['id'])

        #API call using the ID to obtain the cast list from film listed in the text file
        castCompareURL = "https://api.themoviedb.org/3/movie/" + filmCompareID + "/credits?api_key=" + APIKEY

        #Converts JSON in Python readable format
        compareCastResponse = requests.get(castCompareURL)
        compareCastGrab = compareCastResponse.json()

        #Places entire cast list of film from text file into a dynamic list
        for x in compareCastGrab['cast']:
            comparedFilmCastList.append(x['name'])

        #Goes through and grabs the name of the role that the actor/actress portrayed in the film. Only lists one if they have multiple in the same movie
        #It then makes a line for the new text file and puts it into a string array
        for x in set(firstFilmCastList).intersection(comparedFilmCastList):
            actorURL = "http://api.tmdb.org/3/search/person?api_key=" + APIKEY + "&query=" + x

            #Convert the JSON from the call into a usable actor/actress
            actorResponse = requests.get(actorURL)
            actorGrab = actorResponse.json()

            #Grab the first actor/actress ID and then break
            for y in actorGrab['results']:
                actorID = str(y['id'])
                break

            #API call to grab character data based on the actor/actress
            characterURL = "https://api.themoviedb.org/3/person/" + actorID + "/movie_credits?api_key=" + APIKEY

            #Convert the JSON from the call into a usable character
            characterResponse = requests.get(characterURL)
            characterGrab = characterResponse.json()

            #For every actor in the list, grab both the characters that the actor played that warrant inclusion in Full Circle
            for y in characterGrab['cast']:
                if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedFirstFilmTitle and str(y['release_date'][0:4]) == firstFilmYear:
                    firstCharacter = y['character']
                if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedComparedFilmTitle and str(y['release_date'][0:4]) == comparedFilmYear:
                    secondCharacter = y['character']

            #Hokey, don't replicate. Will be fixed later. For both characters, if the string is empty add this error message
            if firstCharacter == "":
                firstCharacter = "**UNABLE TO PULL CHARACTER DATA**"
            if secondCharacter == "":
                secondCharacter = "**UNABLE TO PULL CHARACTER DATA**"
            
            #Formatting message for the Full Circle final output
            finalActorList.append(x + " plays " + firstCharacter + " in " + firstFilmTitle + " (" + firstFilmYear + ") and plays " + secondCharacter + " in " + comparedFilmTitle + "(" + comparedFilmYear + ")")

            firstCharacter = ""     #Set the first film character string back to empty
            secondCharacter = ""    #Set the second film character string back to empty

        #Clears out the compare arrays for the next go around
        comparedFilmCastList.clear()

        #Prevents the while loop from stepping past the length of the file and crashing/ending prematurely
        if iterator + 1 > len(filmCompareList):
            break
        else:
            iterator = iterator + 1

        #Used to reduce the amount of API calls the program makes (otherwise the program might crash)
        time.sleep(0.125)

    #Data for Aladdin (1992)
    comparedFilmTitle = "Aladdin"
    comparedFilmYear = "1992"
    reducedComparedFilmTitle = re.sub('[^A-Za-z0-9]+', '', comparedFilmTitle).lower()
 
    castCompareURL = "https://api.themoviedb.org/3/movie/812/credits?api_key=" + APIKEY

    #Converts JSON in Python readable format
    compareCastResponse = requests.get(castCompareURL)
    compareCastGrab = compareCastResponse.json()

    #Places entire cast list of film from text file into a dynamic list
    for x in compareCastGrab['cast']:
        comparedFilmCastList.append(x['name'])

    #Goes through and grabs the name of the role that the actor/actress portrayed in the film. Only lists one if they have multiple in the same movie
    #It then makes a line for the new text file and puts it into a string array
    for x in set(firstFilmCastList).intersection(comparedFilmCastList):
        actorURL = "http://api.tmdb.org/3/search/person?api_key=" + APIKEY + "&query=" + x

        #Convert the JSON from the call into a usable actor/actress
        actorResponse = requests.get(actorURL)
        actorGrab = actorResponse.json()

        #Grab the first actor/actress ID and then break
        for y in actorGrab['results']:
            actorID = str(y['id'])
            break

        #API call to grab character data based on the actor/actress
        characterURL = "https://api.themoviedb.org/3/person/" + actorID + "/movie_credits?api_key=" + APIKEY

        #Convert the JSON from the call into a usable character
        characterResponse = requests.get(characterURL)
        characterGrab = characterResponse.json()

        #For every actor in the list, grab both the characters that the actor played that warrant inclusion in Full Circle
        for y in characterGrab['cast']:
            if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedFirstFilmTitle and str(y['release_date'][0:4]) == firstFilmYear:
                firstCharacter = y['character']
            if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedComparedFilmTitle and str(y['release_date'][0:4]) == comparedFilmYear:
                secondCharacter = y['character']

        #Hokey, don't replicate. Will be fixed later. For both characters, if the string is empty add this error message
        if firstCharacter == "":
            firstCharacter = "**UNABLE TO PULL CHARACTER DATA**"
        if secondCharacter == "":
            secondCharacter = "**UNABLE TO PULL CHARACTER DATA**"

        #Formatting message for the Full Circle final output
        finalActorList.append(x + " plays " + firstCharacter + " in " + firstFilmTitle + " (" + firstFilmYear + ") and plays " + secondCharacter + " in " + comparedFilmTitle + "(" + comparedFilmYear + ")")

        firstCharacter = ""     #Set the first film character string back to empty
        secondCharacter = ""    #Set the second film character string back to empty

    #Clears out the compare arrays for the next go around
    comparedFilmCastList.clear()

    #Grabs the cast list for It (1990)
    comparedFilmTitle = "It"
    comparedFilmYear = "1990"
    reducedComparedFilmTitle = re.sub('[^A-Za-z0-9]+', '', comparedFilmTitle).lower()

    castCompareURL = "https://api.themoviedb.org/3/tv/19614/credits?api_key=" + APIKEY

    #Converts JSON in Python readable format
    compareCastResponse = requests.get(castCompareURL)
    compareCastGrab = compareCastResponse.json()

    #Places entire cast list of film from text file into a dynamic list
    for x in compareCastGrab['cast']:
        comparedFilmCastList.append(x['name'])

    #Goes through and grabs the name of the role that the actor/actress portrayed in the film. Only lists one if they have multiple in the same movie
    #It then makes a line for the new text file and puts it into a string array
    for x in set(firstFilmCastList).intersection(comparedFilmCastList):
        actorURL = "http://api.tmdb.org/3/search/person?api_key=" + APIKEY + "&query=" + x

        #Convert the JSON from the call into a usable actor/actress
        actorResponse = requests.get(actorURL)
        actorGrab = actorResponse.json()

        #Grab the first actor/actress ID and then break
        for y in actorGrab['results']:
            actorID = str(y['id'])
            break

        #API call to grab character data based on the actor/actress
        characterURL = "https://api.themoviedb.org/3/person/" + actorID + "/movie_credits?api_key=" + APIKEY

        #Convert the JSON from the call into a usable character
        characterResponse = requests.get(characterURL)
        characterGrab = characterResponse.json()

        #API call to grab that same character data, but for their tv roles as It is a miniseries
        characterTVURL = "https://api.themoviedb.org/3/person/" + actorID + "/tv_credits?api_key=" + APIKEY

        #Convert the JSON from the call into a usable character
        characterTVResponse = requests.get(characterTVURL)
        characterTVGrab = characterTVResponse.json()

        #For every actor in the list, grab both the characters that the actor played that warrant inclusion in Full Circle
        for y in characterGrab['cast']:
            if re.sub('[^A-Za-z0-9]+', '', y['title']).lower() == reducedFirstFilmTitle and str(y['release_date'][0:4]) == firstFilmYear:
                firstCharacter = y['character']
        
        #Do the same, but for the character they played in the It miniseries
        for y in characterTVGrab['cast']:
            if re.sub('[^A-Za-z0-9]+', '', y['name']).lower() == reducedComparedFilmTitle:
                secondCharacter = y['character']

        #Hokey, don't replicate. Will be fixed later. For both characters, if the string is empty add this error message
        if firstCharacter == "":
            firstCharacter = "**UNABLE TO PULL CHARACTER DATA**"
        if secondCharacter == "":
            secondCharacter = "**UNABLE TO PULL CHARACTER DATA**"

        #Formatting message for the Full Circle final output
        finalActorList.append(x + " plays " + firstCharacter + " in " + firstFilmTitle + " (" + firstFilmYear + ") and plays " + secondCharacter + " in " + comparedFilmTitle + "(" + comparedFilmYear + ")")

        firstCharacter = ""     #Set the first film character string back to empty
        secondCharacter = ""    #Set the second film character string back to empty

    #Clears out the compare arrays for the next go around
    comparedFilmCastList.clear()

    #Sort the list of actors//actresses alphabetically
    finalActorList.sort()

    #For every value between 0 and the length of the list, this section of code will compare the first sections (with actor name and user's film)
    #of different sections of the list and will condense those attributed to the same actor in order to make things more readable
    for y in range(0, len(finalActorList)):
        if finalActorList[y] != "":
            for z in range (y + 1, len(finalActorList)):
                if finalActorList[y][0:finalActorList[y].index("and plays")] == finalActorList[z][0:finalActorList[z].index("and plays")]:
                    finalActorList[y] = finalActorList[y] + " " + finalActorList[z][finalActorList[z].index("and plays"):len(finalActorList[z])]
                    finalActorList[z] = ""

    return(finalActorList)    #Return the list of Full Circle cases










#Function for creating the show notes document
def theyRemadeItFunc():
    if firstFilmTitleEntry.index("end") == 0 or firstFilmYearEntry.index("end") == 0 or secondFilmTitleEntry.index("end") == 0 or secondFilmYearEntry.index("end") == 0:
        messagebox.showerror("ERROR!", "You're missing some combination of movie titles/years. These are necessary for the program to function")
    elif openLineRadio.get() == 0 and openingLineEntry.index("end") == 0:
        messagebox.showerror("ERROR!", "Your Opening Line is enabled, either ensure it isn't blank or disable it")
    else:
        #This is all for testing whether the film title/year values are valid
        filmID1 = None
        filmID2 = None
        reducedFilmTitle1 = re.sub('[^A-Za-z0-9]+', '', firstFilmTitleEntry.get()).lower()
        filmID_URL = "https://api.themoviedb.org/3/search/movie?api_key=" + APIKEY + "&query=" + firstFilmTitleEntry.get() + "&year=" + firstFilmYearEntry.get()
        userResponse = requests.get(filmID_URL)
        IDGrab = userResponse.json()
        for x in IDGrab['results']:
            if str(x['release_date'][0:4]) == firstFilmYearEntry.get() and re.sub('[^A-Za-z0-9]+', '', str(x['title'])).lower() == reducedFilmTitle1:
                filmID1 = str(x['id'])
        reducedFilmTitle2 = re.sub('[^A-Za-z0-9]+', '', secondFilmTitleEntry.get()).lower()
        filmID_URL = "https://api.themoviedb.org/3/search/movie?api_key=" + APIKEY + "&query=" + secondFilmTitleEntry.get() + "&year=" + secondFilmYearEntry.get()
        userResponse = requests.get(filmID_URL)
        IDGrab = userResponse.json()
        for x in IDGrab['results']:
            if str(x['release_date'][0:4]) == secondFilmYearEntry.get() and re.sub('[^A-Za-z0-9]+', '', str(x['title'])).lower() == reducedFilmTitle2:
                filmID2 = str(x['id'])

        if filmID1 is None or filmID2 is None:
            print(filmID1)
            print(filmID2)
            messagebox.showerror("ERROR!", "Could not find your films, please ensure everything is entered correctly")
        else:
            document = Document()    #Initialize the document
        
            #If the OpenLineRadio button is set to enabled, add in the Opening Line section of the notes
            if openLineRadio.get() == 0:
                document.add_heading('Opening Line:', 1)
                document.add_paragraph(openingLineEntry.get())
            
            if len(listOfWatchedFilms) != 0:
                #Add in the header for the What I've Been Watching section of the notes
                document.add_heading('What I\'ve Been Watching:', 1)
            
                #For every item in the watched films list, add the item to the document under the What I've Been Watching section of the notes
                for y in listOfWatchedFilms:
                    if y != "":
                        document.add_paragraph(y)
        
            #Add in the header for the first Cast List section of the notes
            document.add_heading(firstFilmTitleEntry.get() + ' (' + firstFilmYearEntry.get() + ')' + ' Cast List:', 1)
            listCastList = castListFunc(firstFilmTitleEntry.get(), firstFilmYearEntry.get())    #Set the returned list to the listCastList list
            
            #For every item in the cast list, add the item under the first Cast List section of the notes
            for y in listCastList:
                if y != "":
                    document.add_paragraph(y)
            
            #Add in the header for the second Cast List section of the notes
            document.add_heading(secondFilmTitleEntry.get() + ' (' + secondFilmYearEntry.get() + ')' + ' Cast List:', 1)
            listCastList = castListFunc(secondFilmTitleEntry.get(), secondFilmYearEntry.get())    #Set the returned list to the listCastList list
            
            #For every item in the cast list, add the item under the second Cast List section of the notes
            for y in listCastList:
                if y != "":
                    document.add_paragraph(y)
        
            #Add in the header for the Full Circle section of the notes
            document.add_heading('Full Circle:', 1)
            listFullCircle = fullCircleFunc(firstFilmTitleEntry.get(), firstFilmYearEntry.get())    #Run the Full Circle function for the first film
        
            #For every item in the full circle list, add the item under the Full Circle section of the notes
            for y in listFullCircle:
                if y != "":
                    document.add_paragraph(y)
        
            listFullCircle = fullCircleFunc(secondFilmTitleEntry.get(), secondFilmYearEntry.get())    #Run the Full Circle function for the second film
        
            #For every item in the full circle list, add the item under the Full Circle section of the notes
            for y in listFullCircle:
                if y != "":
                    document.add_paragraph(y)
        
            #If the plotRadio button is enabled, then add in sections for the plots of both films
            if openLineRadio.get() == 0:
                
                #Add heading for the second plot, and then grab the first result for the wikipedia search
                document.add_heading(firstFilmTitleEntry.get() + ' ' + firstFilmYearEntry.get() + ' Plot:', 1)
                wikiResults = wikipedia.search(firstFilmTitleEntry.get() + ' ' + firstFilmYearEntry.get())
                for y in wikiResults:
                    searchPage = str(y)
                    break
        
                #Obtain only the plot section from that page and add it to the document
                listFirstPlot = wikipedia.WikipediaPage(title=searchPage).section("Plot")
                document.add_paragraph(listFirstPlot)
        
                #Add heading for the second plot, and then grab the first result for the wikipedia search
                document.add_heading(secondFilmTitleEntry.get() + ' ' + secondFilmYearEntry.get() + ' Plot:', 1)
                wikiResults = wikipedia.search(secondFilmTitleEntry.get() + ' ' + secondFilmYearEntry.get())
                for y in wikiResults:
                    searchPage = str(y)
                    break
                
                #Obtain only the plot section from that page and add it to the document
                listSecondPlot = wikipedia.WikipediaPage(title=searchPage).section("Plot")
                document.add_paragraph(listSecondPlot)
        
            #Store proposed file title into string var
            titleOfFile = firstFilmTitleEntry.get() + ' (' + firstFilmYearEntry.get() + ') and ' + secondFilmTitleEntry.get() + ' (' + secondFilmYearEntry.get() + ')' + '.docx'
        
            #Save off the document under the freshly created title
            document.save('../../podcastNotes/' + titleOfFile)
        
            #Print out the file that you just created
            if printRadio.get() == 1:
                os.startfile(titleOfFile, "print")
            
            with open("../../textFiles/theyRemadeItList.txt", 'a') as addToFile:
                addToFile.write("\n")
                addToFile.write(firstFilmTitleEntry.get() + " " + firstFilmYearEntry.get())
                addToFile.write("\n")
                addToFile.write(secondFilmTitleEntry.get() + " " + secondFilmYearEntry.get())
                addToFile.close()

            messagebox.showinfo(title="Done", message="Your document has been created and placed in the podcastNotes folder!")

#Application Window
master = Tk()               #Initialization of the application
master.resizable(False, False)
openLineRadio = IntVar()    #Initialization of the Opening Line enable/disable radio button
printRadio = IntVar()       #Initialization of the Print enable/disable radio button
master.title("They Remade It")      #Application title
master.geometry("540x380")          #Application size
labelFont = tkFont.Font(family="Arial", size=14, weight=tkFont.BOLD)     #Styling variables for the fonts of the labels

masterStyle = ttk.Style()                               #Intitialize a Style variable for the application's layout control
masterStyle.configure('TNotebook', tabposition='ne')    #'ne' positions tabs to the NorthEast of the application

tabControl = ttk.Notebook(master)    #Initialize the tabs

tab1 = ttk.Frame(tabControl)                           #Frame for tab 1
tab2 = ttk.Frame(tabControl)                           #Frame for tab 2
tab3 = ttk.Frame(tabControl)                           #Frame for tab 3
tab4 = ttk.Frame(tabControl)                           #Frame for tab 5
tabControl.add(tab1, text='Movies Being Discussed')    #Addition of Movies Being Discussed tab
tabControl.add(tab2, text='Opening Line')              #Addition of Opening Line tab
tabControl.add(tab3, text='Movies Watched')            #Addition of Movies Watched tab
tabControl.add(tab4, text='Finish Up')                 #Addition of Finish Up tab
tabControl.pack(expand=1, fill="both")                 #The "pack" makes the tabs visible in the application

Label(tab1, text="First Film Title:", font=labelFont).place(x=50, y=70)      #Label for the First Film Title text entry field
Label(tab1, text="First Film Year:", font=labelFont).place(x=50, y=120)      #Label for the First Film Year text entry field
Label(tab1, text="Second Film Title:", font=labelFont).place(x=50, y=170)    #Label for the Second Film Title text entry field
Label(tab1, text="Second Film Year:", font=labelFont).place(x=50, y=220)     #Label for the Second Film Year text entry field
Label(tab2, text="Opening Line:", font=labelFont).place(x=50, y=120)         #Label for the Opening Title text entry field
Label(tab3, text="Film Watched:", font=labelFont).place(x=50, y=70)          #Label for the Film Watched text entry field
Label(tab4, text="Final Output:", font=labelFont).place(x=100, y=20)         #Label for the Final Read Only text output field
Label(tab4, text="Print?", font=labelFont).place(x=300, y=125)               #Label for the Print options

firstFilmTitleEntry = Entry(tab1)     #Text entry field for the First Film Title
firstFilmYearEntry = Entry(tab1)      #Text entry field for the First Film Year
secondFilmTitleEntry = Entry(tab1)    #Text entry field for the Second Film Title
secondFilmYearEntry = Entry(tab1)     #Text entry field for the Second Film Year
openingLineEntry = Entry(tab2)        #Text entry field for the Opening Title
filmWatchedEntry = Entry(tab3)        #Text entry field for Film Watched
filmWatchText = Text(tab3)            #Text entry for Film Watched
filmWatchText.configure(state="disabled")    #Set the Film Watched textbox to disabled (for read-only purposes)
readOnlyText = Text(tab4)             #Text field for Final Output
readOnlyText.tag_configure("bold", font="Helvetica 12 bold")   #
scrollbarFinal = Scrollbar(master)          #Initialization of scrollbar
scrollbarWatched = Scrollbar(master)        #Initialization of scrollbar
scrollbarFinal.pack(side=RIGHT, fill=Y)       #Packing the scrollbar to be on the righthand side of the textbox its assigned to
scrollbarWatched.pack(side=RIGHT, fill=Y)     #Packing the scrollbar to be on the righthand side of the textbox its assigned to
readOnlyText.config(yscrollcommand=scrollbarFinal.set)         #Set the films watched textbox to be controlled by the scrollbar
filmWatchText.config(yscrollcommand=scrollbarWatched.set)      #Set the final output textbox to be controlled by the scrollbar
scrollbarFinal.config(command=readOnlyText.yview)              #Configure the scrollbar to control the y axis for the films watched textbox
scrollbarWatched.config(command=filmWatchText.yview)           #Configure the scrollbar to control the y axis for the final output textbox
scrollbarFinal.place(in_=readOnlyText, relx=1.0, relheight=1.0, bordermode="outside")         #Placement for the scrollbar
scrollbarWatched.place(in_=filmWatchText, relx=1.0, relheight=1.0, bordermode="outside")      #Placement for the scrollbar

openingLineEntry.place(x=250, y=125, height=20, width=250)        #Placement for the Opening Title text entry field
filmWatchedEntry.place(x=250, y=75, height=20, width=250)         #Placement for the Film Watched text entry field
filmWatchText.place(x=35, y=170, height=150, width=470)           #Placement for the Final Output text field
firstFilmTitleEntry.place(x=250, y=75, height=20, width=250)      #Placement for the First Film Title text entry field
firstFilmYearEntry.place(x=250, y=125, height=20, width=40)       #Placement for the First Film Year text entry field
secondFilmTitleEntry.place(x=250, y=175, height=20, width=250)    #Placement for the Second Film Title text entry field
secondFilmYearEntry.place(x=250, y=225, height=20, width=40)      #Placement for the Second Film Year text entry field
readOnlyText.place(x=35, y=50, height=275, width=250)             #Placement for the Final Output text field

#Setting up the enable button for the Opening Line Widget
enableOpeningLineEntry = Radiobutton(tab2, text="Yes", variable=openLineRadio, value="0", command=enableOpeningLineEntryFunc)
enableOpeningLineEntry.place(x=50, y=170)     #Placement for the Opening Line enable button

#Setting up the disable button for the Opening Line Widget
disableOpeningLineEntry = Radiobutton(tab2, text="No", variable=openLineRadio, value="1", command=disableOpeningLineEntryFunc)
disableOpeningLineEntry.place(x=150, y=170)   #Placement for the Opening Line disable button

#Setting up the enable button for printing
enablePrint = Radiobutton(tab4, text="Enable", variable=printRadio, value="1")
enablePrint.place(x=375, y=128)    #Placement for the Print enable button

#Setting up the disable button for printing
disablePrint = Radiobutton(tab4, text="Disable", variable=printRadio, value="0")
disablePrint.place(x=450, y=128)    #Placement for the Print disable button

#Setting up the Add Film button, for putting films into the Watched list
Button(tab3, text='Add Film', command=addToWatchListFunc).place(x=52, y=120)

#Setting up the Submit button for creating the show notes doc
Button(tab4, text='Submit', command=theyRemadeItFunc).place(x=385, y=160)

#Calls the Function to populate the read only textbox whenever a new tab is selected. This is so that the field will update consistently
tabControl.bind("<<NotebookTabChanged>>", onTabClickFunc)

master.mainloop()    #Keeps the Application open until the user closes it